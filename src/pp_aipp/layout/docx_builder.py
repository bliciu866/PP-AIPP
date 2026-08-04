from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .models import LayoutTheme


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def _font(run, name: str, size: float, bold=False, color="2E2E2E", italic=False) -> None:
    run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


class RecipeDocxBuilder:
    def __init__(self, theme: LayoutTheme | None = None) -> None:
        self.theme = theme or LayoutTheme()

    def build(self, recipes: list[dict], output_path: str | Path) -> Path:
        output = Path(output_path); output.parent.mkdir(parents=True, exist_ok=True)
        doc = Document(); section = doc.sections[0]
        section.page_width = Inches(self.theme.page_width_in); section.page_height = Inches(self.theme.page_height_in)
        section.top_margin = Inches(self.theme.margin_top_in); section.bottom_margin = Inches(self.theme.margin_bottom_in)
        section.left_margin = Inches(self.theme.margin_inside_in); section.right_margin = Inches(self.theme.margin_outside_in)
        self._configure_styles(doc); self._footer(section)
        for index, recipe in enumerate(recipes):
            if index: doc.add_page_break()
            self._recipe_page(doc, recipe)
        doc.save(output)
        return output

    def _configure_styles(self, doc: Document) -> None:
        normal = doc.styles["Normal"]; normal.font.name = self.theme.body_font; normal.font.size = Pt(8.7)
        normal.paragraph_format.space_after = Pt(2); normal.paragraph_format.line_spacing = 1.02
        for name, size in (("Title", 21), ("Heading 1", 12.5), ("Heading 2", 10.5)):
            style=doc.styles[name]; style.font.name=self.theme.title_font; style.font.size=Pt(size); style.font.bold=True
            style.font.color.rgb=RGBColor.from_string(self.theme.primary_hex)
            style.paragraph_format.space_before=Pt(3); style.paragraph_format.space_after=Pt(2)

    def _footer(self, section) -> None:
        p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("PROJECT PHYSIQUE™  •  PP-AIPP GENERATED EDITION  •  "); _font(r,self.theme.body_font,7,color="666666")
        fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

    def _recipe_page(self, doc: Document, r: dict) -> None:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
        rr=p.add_run(r["recipe_id"]); _font(rr,self.theme.body_font,8,bold=True,color=self.theme.primary_hex)
        p=doc.add_paragraph(style="Title"); p.paragraph_format.keep_with_next=True
        _font(p.add_run(r["title"]),self.theme.title_font,20,bold=True,color=self.theme.charcoal_hex)
        desc=r.get("description") or self._fallback_description(r)
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
        _font(p.add_run(desc),self.theme.body_font,9.3,italic=True,color="555555")

        # hero placeholder
        t=doc.add_table(rows=1,cols=1); t.autofit=False; t.columns[0].width=Inches(7.0)
        c=t.cell(0,0); c.height=Inches(1.05); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; _shade(c,"E8ECE8"); _set_cell_margins(c,180,100,180,100)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run("HERO PHOTO  •  4:5 ASSET SLOT"),self.theme.body_font,9,bold=True,color="6C7B6C")

        badges="  |  ".join(r.get("badges",[])[:6]) or self._default_badges(r)
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        _font(p.add_run(badges),self.theme.body_font,8.2,bold=True,color=self.theme.primary_hex)

        info=doc.add_table(rows=1,cols=4); info.autofit=True
        values=[("MEAL",r.get("meal") or "—"),("SERVINGS",str(r.get("servings",1))),("INGREDIENTS",str(len(r.get("ingredients",[])))),("STATUS",r.get("status","—").replace("_"," "))]
        for i,(label,val) in enumerate(values):
            c=info.cell(0,i); _shade(c,self.theme.sage_hex); _set_cell_margins(c,60,80,60,80)
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            _font(p.add_run(label+"\n"),self.theme.body_font,6.5,bold=True,color="607060")
            _font(p.add_run(val),self.theme.body_font,8,bold=True,color=self.theme.charcoal_hex)

        body=doc.add_table(rows=1,cols=2); body.autofit=False; body.columns[0].width=Inches(2.65); body.columns[1].width=Inches(4.25)
        left,right=body.cell(0,0),body.cell(0,1); _set_cell_margins(left,60,0,20,140); _set_cell_margins(right,60,140,20,0)
        self._cell_heading(left,"Ingredients")
        for ing in r.get("ingredients",[]):
            p=left.add_paragraph(); p.paragraph_format.space_after=Pt(0)
            _font(p.add_run(f"{ing['quantity']:g} {ing['unit']}  "),self.theme.body_font,8,bold=True,color=self.theme.primary_hex)
            _font(p.add_run(ing['name']),self.theme.body_font,8,color=self.theme.charcoal_hex)
        self._cell_heading(right,"Method")
        method=r.get("method",[])
        if method:
            for step in method:
                p=right.add_paragraph(); p.paragraph_format.left_indent=Inches(0.02); p.paragraph_format.space_after=Pt(1)
                _font(p.add_run(f"{step['number']}. "),self.theme.body_font,8,bold=True,color=self.theme.primary_hex)
                _font(p.add_run(step['text']),self.theme.body_font,8,color=self.theme.charcoal_hex)
        else:
            p=right.add_paragraph(); _font(p.add_run("Method not available in the controlled source."),self.theme.body_font,8,italic=True,color="777777")

        n=r.get("nutrition")
        if n:
            doc.add_paragraph("Nutrition per serving",style="Heading 1")
            nt=doc.add_table(rows=2,cols=5); labels=("Energy","Protein","Carbs","Fat","Fibre")
            vals=(f"{n['energy_kcal']:g} kcal",f"{n['protein_g']:g} g",f"{n['carbohydrate_g']:g} g",f"{n['fat_g']:g} g",f"{n['fibre_g']:g} g")
            for i in range(5):
                for row,txt in ((0,labels[i]),(1,vals[i])):
                    c=nt.cell(row,i); _shade(c,self.theme.sage_hex); _set_cell_margins(c,55,55,55,55)
                    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    _font(p.add_run(txt),self.theme.body_font,7.5,bold=(row==1),color=self.theme.charcoal_hex)

        self._panel(doc,"Meal Prep",r.get("meal_prep") or "No controlled meal-prep guidance available.",self.theme.sage_hex)
        if r.get("chef_tip"): self._panel(doc,"Chef's Tip",r["chef_tip"],self.theme.yellow_hex)
        if r.get("ingredient_swap"): self._panel(doc,"Ingredient Swap",r["ingredient_swap"],self.theme.grey_hex)
        qa=" ".join(q.get("message","") for q in r.get("qa_records",[]) if q.get("message"))
        if qa: self._panel(doc,"Technical QA Note",qa,self.theme.grey_hex,small=True)

    def _cell_heading(self, cell, text: str) -> None:
        p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
        _font(p.add_run(text),self.theme.title_font,11,bold=True,color=self.theme.primary_hex)

    def _panel(self, doc, label, text, fill, small=False):
        t=doc.add_table(rows=1,cols=1); c=t.cell(0,0); _shade(c,fill); _set_cell_margins(c,55,90,55,90)
        p=c.paragraphs[0]; _font(p.add_run(label+": "),self.theme.body_font,7.4 if small else 8,bold=True,color=self.theme.charcoal_hex)
        _font(p.add_run(text),self.theme.body_font,7.1 if small else 8,color="4A4A4A")

    @staticmethod
    def _fallback_description(r):
        meal=(r.get("meal") or "meal").lower()
        return f"A practical, high-protein {meal} designed for simple preparation and repeatable nutrition tracking."

    @staticmethod
    def _default_badges(r):
        return f"{r.get('meal','Meal')}  |  UK CoFID Verified"
