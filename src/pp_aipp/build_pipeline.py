"""Production build pipeline connecting Gold Master import to the layout engine."""
from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from .domain import ProjectDatabase
from .layout import LayoutBuildResult, LayoutEngine
from .parser import GoldMasterImportService, ImportSummary


@dataclass(frozen=True, slots=True)
class BookBuildPipelineResult:
    book_id: str
    database_path: Path
    import_report_path: Path
    layout_report_path: Path
    import_summary: ImportSummary
    layout: LayoutBuildResult


def _book_id(project_root: Path) -> str:
    value = re.sub(r"[^a-z0-9]+", "-", project_root.name.lower()).strip("-")
    return value or "project-physique-book"


def _is_premium_schema(source: Path) -> bool:
    """Return True only for the controlled v5 premium editorial schema."""
    document = Document(source)
    fragments = [paragraph.text for paragraph in document.paragraphs]
    fragments.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    text = "\n".join(fragments).upper()
    return all(
        marker in text
        for marker in (
            "CHEFIE’S TIP",
            "COMMON MISTAKE",
            "INGREDIENT SWAP",
            "SERVING SUGGESTION",
            "30-DAY SUCCESS GUIDE",
            "UK SHOPPING SYSTEM",
            "30-DAY PROGRESS TRACKER",
        )
    )


def build_gold_master_book(
    project_root: str | Path,
    source_docx: str | Path,
    *,
    strict_collection: bool = True,
) -> BookBuildPipelineResult:
    """Parse a controlled DOCX, persist recipes, and generate a built book DOCX."""
    root = Path(project_root).expanduser().resolve()
    source = Path(source_docx).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(source)

    data_dir = root / "data"
    qa_dir = root / "qa"
    build_dir = root / "build"
    exports_dir = root / "exports"
    for directory in (data_dir, qa_dir, build_dir, exports_dir):
        directory.mkdir(parents=True, exist_ok=True)

    book_id = _book_id(root)
    database_path = data_dir / "project.sqlite3"
    import_report = qa_dir / "gold_master_import_report.json"
    layout_report = qa_dir / "layout_build_report.json"
    output_docx = build_dir / "Project_Physique_30_Days_Fat_Loss_Built.docx"

    database = ProjectDatabase(database_path)
    importer = GoldMasterImportService(database)
    summary, parsed = importer.import_docx(
        source,
        book_id=book_id,
        replace=True,
        report_path=import_report,
        strict_collection=strict_collection,
    )
    if parsed.issues:
        errors = [issue for issue in parsed.issues if issue.severity == "ERROR"]
        if errors:
            detail = "; ".join(f"{issue.code}: {issue.message}" for issue in errors[:5])
            raise ValueError(f"Gold Master parse failed: {detail}")
    if summary.imported_recipes == 0:
        raise ValueError("Gold Master contains no importable recipes")

    if _is_premium_schema(source):
        # Premium passthrough mode is deliberate: the source already contains
        # publication layout, front/back matter and recipe callouts that the
        # legacy database renderer cannot represent without data loss.
        shutil.copy2(source, output_docx)
        preview = source.with_suffix(".preview.pdf")
        output_pdf = output_docx.with_suffix(".pdf")
        if preview.is_file():
            shutil.copy2(preview, output_pdf)
        else:
            output_pdf = None
        layout = LayoutBuildResult(
            output_docx=output_docx,
            output_pdf=output_pdf,
            recipe_count=summary.imported_recipes,
            page_breaks=0,
            warnings=["PREMIUM_SCHEMA_PASSTHROUGH: complete editorial layout preserved"],
        )
    else:
        layout = LayoutEngine(database_path).build_book(output_docx, book_id=book_id, pdf=False)
    LayoutEngine.write_report(layout, layout_report)
    return BookBuildPipelineResult(
        book_id=book_id,
        database_path=database_path,
        import_report_path=import_report,
        layout_report_path=layout_report,
        import_summary=summary,
        layout=layout,
    )
