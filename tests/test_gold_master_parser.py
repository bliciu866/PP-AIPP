from pathlib import Path

import pytest
from docx import Document

from pp_aipp.parser import GoldMasterParser


def build_sample(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("PP-R001")
    doc.add_paragraph("Sample Protein Breakfast")
    doc.add_paragraph("🥣 Breakfast | 💪 High Protein | ✓ UK CoFID Verified")
    info = doc.add_table(rows=2, cols=4)
    info.rows[0].cells[0].text = "Meal"
    info.rows[0].cells[1].text = "Breakfast"
    info.rows[0].cells[2].text = "Servings"
    info.rows[0].cells[3].text = "1"
    info.rows[1].cells[0].text = "Ingredients"
    info.rows[1].cells[1].text = "2"
    info.rows[1].cells[2].text = "Status"
    info.rows[1].cells[3].text = "NUTRITION LOCKED"
    doc.add_paragraph("Ingredients")
    ing = doc.add_table(rows=3, cols=2)
    ing.rows[0].cells[0].text = "Ingredient"
    ing.rows[0].cells[1].text = "Quantity"
    ing.rows[1].cells[0].text = "Quark"
    ing.rows[1].cells[1].text = "250 g"
    ing.rows[2].cells[0].text = "Blueberries"
    ing.rows[2].cells[1].text = "80 g"
    doc.add_paragraph("Method")
    doc.add_paragraph("1. Mix and serve.")
    doc.add_paragraph("Nutrition per serving")
    nut = doc.add_table(rows=2, cols=5)
    for i, label in enumerate(("Energy", "Protein", "Carbohydrates", "Fat", "Fibre")):
        nut.rows[0].cells[i].text = label
    for i, value in enumerate(("350 kcal", "35 g", "30 g", "8 g", "5 g")):
        nut.rows[1].cells[i].text = value
    doc.add_paragraph("QA note: direct records used.")
    doc.save(path)


def test_parser_creates_domain_recipe(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    build_sample(source)
    result = GoldMasterParser().parse(source, book_id="book-1")
    assert len(result.recipes) == 1
    recipe = result.recipes[0]
    assert recipe.recipe_id == "PP-R001"
    assert recipe.meal == "Breakfast"
    assert len(recipe.ingredients) == 2
    assert recipe.method[0].number == 1
    assert recipe.nutrition and recipe.nutrition.energy_kcal == 350


def test_parser_rejects_missing_file() -> None:
    with pytest.raises(ValueError):
        GoldMasterParser().parse("missing.docx", book_id="book-1")


def test_parser_reads_inline_and_table_controlled_content(tmp_path: Path) -> None:
    source = tmp_path / "content-sections.docx"
    build_sample(source)
    document = Document(source)
    # Add a second controlled recipe with inline and table-based content.
    document.add_paragraph("PP-R002")
    document.add_paragraph("Table Method Recipe")
    document.add_paragraph("Dinner | High Protein | UK CoFID Verified")
    ingredients = document.add_table(rows=2, cols=2)
    ingredients.cell(0, 0).text = "Ingredient"
    ingredients.cell(0, 1).text = "Quantity"
    ingredients.cell(1, 0).text = "Chicken breast"
    ingredients.cell(1, 1).text = "180 g"
    content = document.add_table(rows=2, cols=2)
    content.cell(0, 0).text = "Method"
    content.cell(0, 1).text = "1. Season the chicken.\n2. Roast until cooked through."
    content.cell(1, 0).text = "Meal Prep"
    content.cell(1, 1).text = "Chill and refrigerate for up to 3 days."
    document.add_paragraph("Nutrition per serving")
    nutrition = document.add_table(rows=2, cols=5)
    for i, label in enumerate(("Energy", "Protein", "Carbs", "Fat", "Fibre")):
        nutrition.cell(0, i).text = label
    for i, value in enumerate(("450 kcal", "45 g", "35 g", "12 g", "6 g")):
        nutrition.cell(1, i).text = value
    document.add_paragraph("QA Note: Controlled table content.")
    document.save(source)

    result = GoldMasterParser().parse(source, book_id="book-1", strict_collection=False)
    recipe = next(recipe for recipe in result.recipes if recipe.recipe_id == "PP-R002")
    assert [step.text for step in recipe.method] == [
        "Season the chicken.",
        "Roast until cooked through.",
    ]
    assert recipe.meal_prep == "Chill and refrigerate for up to 3 days."
