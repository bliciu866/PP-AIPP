from pathlib import Path

from docx import Document

from pp_aipp.build_pipeline import build_gold_master_book


def _source(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("PP-R001")
    doc.add_paragraph("Pipeline Test Recipe")
    doc.add_paragraph("A controlled recipe description that is long enough for parser recognition.")
    doc.add_paragraph("Breakfast | High Protein | UK CoFID Verified")
    info = doc.add_table(rows=1, cols=1)
    info.cell(0, 0).text = "Meal Breakfast Servings 1 Status NUTRITION LOCKED"
    doc.add_paragraph("Ingredients")
    ingredients = doc.add_table(rows=2, cols=2)
    ingredients.cell(0, 0).text = "Ingredient"
    ingredients.cell(0, 1).text = "Quantity"
    ingredients.cell(1, 0).text = "Porridge oats"
    ingredients.cell(1, 1).text = "60 g"
    doc.add_paragraph("Method")
    doc.add_paragraph("Mix and serve.")
    doc.add_paragraph("Nutrition per serving")
    nutrition = doc.add_table(rows=2, cols=5)
    for index, value in enumerate(("Energy", "Protein", "Carbs", "Fat", "Fibre")):
        nutrition.cell(0, index).text = value
    for index, value in enumerate(("400 kcal", "30 g", "45 g", "10 g", "7 g")):
        nutrition.cell(1, index).text = value
    doc.add_paragraph("QA Note: Controlled test record.")
    doc.save(path)
    return path


def test_pipeline_builds_docx_database_and_reports(tmp_path):
    project = tmp_path / "30_Days_Fat_Loss"
    source = _source(tmp_path / "gold-master.docx")

    result = build_gold_master_book(project, source, strict_collection=False)

    assert result.import_summary.imported_recipes == 1
    assert result.layout.recipe_count == 1
    assert result.layout.output_docx.is_file()
    assert result.layout.output_docx.parent == project / "build"
    assert result.database_path.is_file()
    assert result.import_report_path.is_file()
    assert result.layout_report_path.is_file()


def test_pipeline_preserves_premium_schema_and_verified_pdf(tmp_path):
    project = tmp_path / "premium"
    source = _source(tmp_path / "premium.docx")
    document = Document(source)
    for marker in (
        "CHEFIE’S TIP",
        "COMMON MISTAKE",
        "INGREDIENT SWAP",
        "SERVING SUGGESTION",
        "Your 30-Day Success Guide",
        "UK Shopping System",
        "30-Day Progress Tracker",
    ):
        document.add_paragraph(marker)
    document.save(source)
    preview = source.with_suffix(".preview.pdf")
    preview.write_bytes(b"%PDF-1.4 premium\n")

    result = build_gold_master_book(project, source, strict_collection=False)

    assert result.layout.output_docx.read_bytes() == source.read_bytes()
    assert result.layout.output_pdf.read_bytes() == preview.read_bytes()
    assert "PREMIUM_SCHEMA_PASSTHROUGH" in result.layout.warnings[0]
